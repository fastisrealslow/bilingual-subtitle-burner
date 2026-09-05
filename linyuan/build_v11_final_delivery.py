#!/usr/bin/env python3
"""Build the final V11 14-part acceptance report and review-asset package."""
import argparse
import json
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SLUG = "ly-parity-v3-14-0905"
RUNS = {
    "render": "33960957254",
    "recovery": "33963713427",
    "publish": "33985111896",
}
GATES = [
    "quality_gate_version",
    "review_assets_verified",
    "cover_person_image_verified",
    "title_quality_verified",
    "no_qr_verified",
    "no_black_bars_verified",
    "watermark_verified",
    "live_region_verified",
]


def set_cell_shading(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shade = props.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        props.append(shade)
    shade.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    props = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    props.append(repeat)


def hyperlink(paragraph, text, url):
    part = paragraph.part
    rel = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1565C0")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    link.append(run)
    paragraph._p.append(link)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color in [
        ("Title", 24, "16324F"),
        ("Heading 1", 16, "16324F"),
        ("Heading 2", 12, "1565C0"),
    ]:
        style = styles[name]
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    styles["Title"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.space_before = Pt(10)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.space_before = Pt(7)
    styles["Heading 2"].paragraph_format.space_after = Pt(3)


def copy_evidence(parts_root, out):
    out.mkdir(parents=True, exist_ok=True)
    metas = None
    for index in range(1, 15):
        src = parts_root / str(index)
        for pattern in [
            f"cover_{index}.jpg",
            f"preview_30s_{index}.mp4",
            f"contact_sheet_6_{index}.jpg",
        ]:
            path = src / pattern
            if not path.is_file():
                raise SystemExit(f"missing evidence: {path}")
            shutil.copy2(path, out / path.name)
        if metas is None:
            candidate = src / "meta.json"
            metas = json.loads(candidate.read_text(encoding="utf-8"))
    if not isinstance(metas, list) or len(metas) != 14:
        raise SystemExit("metadata is not an exact 14-part list")
    (out / "meta.json").write_text(json.dumps(metas, ensure_ascii=False, indent=2), encoding="utf-8")
    return metas


def validate(metas, parts):
    if len(parts) != 14 or len({x.get("bvid") for x in parts}) != 14:
        raise SystemExit("Bilibili receipts are incomplete or duplicated")
    results = []
    for index, item in enumerate(metas, 1):
        checks = {
            "quality_gate_v11": int(item.get("quality_gate_version") or 0) >= 11,
            "review_assets_verified": item.get("review_assets_verified") is True,
            "cover_person_image_verified": item.get("cover_person_image_verified") is True,
            "title_quality_verified": item.get("title_quality_verified") is True,
            "no_qr_verified": item.get("no_qr_verified") is True,
            "no_black_bars_verified": item.get("no_black_bars_verified") is True,
            "watermark_verified": item.get("watermark_verified") is True,
            "live_region_verified": item.get("live_region_verified") is True,
        }
        if not all(checks.values()):
            raise SystemExit(f"part {index} gate failure: {checks}")
        results.append(checks)
    return results


def build_report(path, evidence, metas, receipts, checks):
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("林园 V11 · 14 条出片与投稿最终验收")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run("批次：ly-parity-v3-14-0905  |  完成日期：2026-09-06（北京时间）")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("52606D")

    callout = doc.add_table(rows=1, cols=4)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    widths = [1.75, 1.75, 1.75, 1.75]
    labels = [("成片", "14 / 14"), ("V11 门禁", "14 / 14"), ("投稿回执", "14 / 14"), ("BV 唯一性", "14 / 14")]
    for i, (label, value) in enumerate(labels):
        cell = callout.cell(0, i)
        cell.width = Inches(widths[i])
        set_cell_shading(cell, "EAF2F8")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = p.add_run(value + "\n")
        a.bold = True
        a.font.size = Pt(15)
        a.font.color.rgb = RGBColor.from_string("1565C0")
        b = p.add_run(label)
        b.font.size = Pt(8.5)
        b.font.color.rgb = RGBColor.from_string("52606D")

    doc.add_heading("1. 验收结论", level=1)
    p = doc.add_paragraph()
    p.add_run("结论：").bold = True
    p.add_run("本批次 13 条观点切片与 1 条 57 分钟完整版均已生成，全部通过 V11 质量门禁，并由境内 FC 完成正式投稿。14 个 BV 号均已写入持久状态，数量与唯一性复核通过。")
    p = doc.add_paragraph()
    p.add_run("线上标准：").bold = True
    p.add_run("后续中文源出片默认沿用 V4 真人封面、V11 质量门禁、真人动态版式、无黑色占位、无原视频 URL 的投稿边界；B 站提交使用 Web 通道，规避旧 APP 接口 21566 风控。")

    doc.add_heading("2. 14 条正式投稿回执", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    col_widths = [0.42, 1.35, 4.85, 0.72]
    headers = ["#","BV 号","标题","门禁"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        set_cell_shading(cell, "16324F")
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])
    for i, (receipt, gate) in enumerate(zip(receipts, checks), 1):
        cells = table.add_row().cells
        for j, width in enumerate(col_widths):
            cells[j].width = Inches(width)
        cells[0].paragraphs[0].add_run(str(i))
        hyperlink(cells[1].paragraphs[0], receipt["bvid"], "https://www.bilibili.com/video/" + receipt["bvid"])
        cells[2].paragraphs[0].add_run(receipt["title"])
        ok = cells[3].paragraphs[0].add_run("V11 通过")
        ok.bold = True
        ok.font.color.rgb = RGBColor.from_string("1B7F5A")
        if i % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "F6F8FA")
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.4)

    doc.add_heading("3. 质量门禁明细", level=1)
    doc.add_paragraph("14 条均通过以下 8 项硬门槛：")
    gates = [
        "质量门禁版本 ≥ V11",
        "30 秒样片与 6 帧检查图已生成并校验",
        "封面人物图为已核验林园真人图",
        "标题原话与污染词校验通过",
        "二维码检测通过",
        "黑边/黑色人物占位检测通过",
        "源水印清理与品牌水印检测通过",
        "真人动态有效区域检测通过",
    ]
    for item in gates:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. 实际视觉证据", level=1)
    doc.add_paragraph("以下为第 1 条观点切片与第 14 条完整版的实际 V4 封面代表；完整 14 套封面、30 秒样片和 6 帧检查图均收录在随附证据包。")
    images = doc.add_table(rows=1, cols=2)
    images.alignment = WD_TABLE_ALIGNMENT.CENTER
    images.autofit = False
    for col, index in enumerate([1,14]):
        cell = images.cell(0,col)
        cell.width = Inches(3.45)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(evidence / f"cover_{index}.jpg"), width=Inches(3.15))
        q = cell.add_paragraph(f"第 {index} 条实际封面")
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.runs[0].font.size = Pt(8.5)
        q.runs[0].font.color.rgb = RGBColor.from_string("52606D")

    doc.add_heading("5. 可追溯记录", level=1)
    records = [
        ("V11 全量渲染", RUNS["render"], "14 条渲染完成；聚合包仅在 FinalizeArtifact 超时"),
        ("逐条恢复与复核", RUNS["recovery"], "14 个独立快通道包；14 条真实 verify_render 与门禁证明"),
        ("最终异步 Web 投稿", RUNS["publish"], "第 14 条完整版异步提交；最终 14/14"),
    ]
    rec = doc.add_table(rows=1, cols=3)
    rec.style = "Table Grid"
    rec.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,text in enumerate(["环节","运行编号","结果"]):
        cell=rec.rows[0].cells[i]
        set_cell_shading(cell,"DDEBF7")
        cell.paragraphs[0].add_run(text).bold=True
    for name,run_id,result in records:
        cells=rec.add_row().cells
        cells[0].text=name
        cells[1].text=run_id
        cells[2].text=result

    doc.add_heading("6. 本轮修复与后续默认规则", level=1)
    fixes = [
        "横竖封面标题容量统一为 14 字 × 3 行，避免第 3 条 42 字标题被误拦。",
        "人物核验采用六帧首检与分组三帧复核；人物图提取失败直接终止，不再使用黑色占位。",
        "投稿前再次清洗简介 URL；转载来源仅传平台文字，不传原 B 站地址。",
        "精确批次恢复参数固定为 13 条切片 + 1 条完整版，禁止退化为 3 条。",
        "大文件按逐条快通道恢复；57 分钟完整版通过 FC 异步调用，避免同步网关 85 秒上限。",
        "正式提交改用 Biliup 已合并的 Web 接口，解决旧 APP 接口 21566 风控。",
    ]
    for item in fixes:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. 证据包目录", level=1)
    doc.add_paragraph("随附 ZIP 包含：")
    for item in [
        "cover_1.jpg 至 cover_14.jpg（14 张实际封面）",
        "preview_30s_1.mp4 至 preview_30s_14.mp4（14 个 30 秒样片）",
        "contact_sheet_6_1.jpg 至 contact_sheet_6_14.jpg（14 张六帧检查图）",
        "meta.json（14 条完整 V11 门禁、版式、指纹与素材证明）",
        "receipts.json（14 条标题、BV 号与 B 站链接）",
        "quality_summary.json（14 × 8 项质量门禁核验结果）",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("林园 V11 14 条最终验收 · ly-parity-v3-14-0905").font.size = Pt(8)
    doc.save(path)


def main():
    parser=argparse.ArgumentParser()
    parser.add_argument("--parts",required=True)
    parser.add_argument("--state",required=True)
    parser.add_argument("--out",required=True)
    args=parser.parse_args()
    root=Path(args.out)
    evidence=root/"evidence"
    metas=copy_evidence(Path(args.parts),evidence)
    state=json.loads(Path(args.state).read_text(encoding="utf-8"))
    pub=state.get("published",{}).get(SLUG,{})
    receipts=[x for x in pub.get("parts",[]) if x.get("status")=="published"]
    checks=validate(metas,receipts)
    (evidence/"receipts.json").write_text(json.dumps(receipts,ensure_ascii=False,indent=2),encoding="utf-8")
    (evidence/"quality_summary.json").write_text(json.dumps(checks,ensure_ascii=False,indent=2),encoding="utf-8")
    report=root/"林园V11-14条出片投稿最终验收报告.docx"
    build_report(report,evidence,metas,receipts,checks)
    archive=root/"林园V11-14条封面样片检查图证据包.zip"
    with zipfile.ZipFile(archive,"w",zipfile.ZIP_DEFLATED,compresslevel=6) as z:
        for path in sorted(evidence.rglob("*")):
            if path.is_file():
                z.write(path,path.relative_to(evidence))
    print(json.dumps({"report":str(report),"evidence":str(archive),"receipts":len(receipts),"gates":len(checks)},ensure_ascii=False))


if __name__=="__main__":
    main()
